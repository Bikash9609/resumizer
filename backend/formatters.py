from markdown_pdf import MarkdownPdf, Section
from docx import Document
import io
import tempfile
import os


def generate_pdf_from_markdown(
    markdown_content: str, template_type: str = "standard"
) -> io.BytesIO:
    """Creates a PDF from a Markdown string."""
    styles = ""
    if template_type == "modern":
        styles = "<style>body { font-family: 'Helvetica Neue', Arial, sans-serif; color: #333; line-height: 1.6; } h1 { color: #2c3e50; border-bottom: 2px solid #3498db; } h2 { color: #2980b9; }</style>\n\n"
    elif template_type == "compact":
        styles = "<style>body { font-family: 'Arial', sans-serif; line-height: 1.2; font-size: 11pt; } h1, h2, h3 { margin-bottom: 2px; margin-top: 5px; } p, ul { margin-bottom: 5px; margin-top: 5px; }</style>\n\n"
    elif template_type == "creative":
        styles = "<style>body { font-family: 'Georgia', serif; color: #444; } h1 { color: #e74c3c; text-align: center; font-size: 2.2em; text-transform: uppercase; letter-spacing: 2px; } h2 { color: #c0392b; border-bottom: 1px dashed #e74c3c; }</style>\n\n"
    elif template_type == "executive":
        styles = "<style>body { font-family: 'Palatino Linotype', 'Book Antiqua', Palatino, serif; line-height: 1.5; color: #222; } h1 { font-size: 2em; text-align: center; letter-spacing: 1px; } h2 { border-bottom: 1px solid #000; text-transform: uppercase; font-size: 1.2em; }</style>\n\n"
    else:  # standard
        styles = "<style>body { font-family: 'Times New Roman', Times, serif; line-height: 1.4; color: #000; } h1, h2 { border-bottom: 1px solid #ccc; }</style>\n\n"

    styled_content = f"{styles}{markdown_content}"

    pdf = MarkdownPdf(toc_level=2)
    pdf.add_section(Section(styled_content))

    # Save to a bytes buffer
    buffer = io.BytesIO()
    # The markdown-pdf library usually writes to a file path
    # A simple trick is to write it to a temp file and read it back, but let's try writing to buffer if supported, or temp file.

    with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp:
        temp_path = tmp.name

    pdf.save(temp_path)

    with open(temp_path, "rb") as f:
        buffer.write(f.read())

    os.remove(temp_path)
    buffer.seek(0)
    return buffer


def generate_docx_from_markdown(markdown_content: str) -> io.BytesIO:
    """Creates a very basic DOCX from Markdown (simplified for prototype)."""
    document = Document()
    document.add_heading("Generated Resume", 0)

    # Extremely simplified parsing - for a real app, use a proper md to docx library like pypandoc
    for line in markdown_content.split("\n"):
        if line.startswith("#"):
            level = line.count("#")
            text = line.replace("#", "").strip()
            document.add_heading(text, level=min(level, 9))
        else:
            document.add_paragraph(line)

    buffer = io.BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer
